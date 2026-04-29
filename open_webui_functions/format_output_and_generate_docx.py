"""
title: Test Case JSON → Table + DOCX Generator
author: Rishikesh Kumar
version: 1.0
"""

import json
import re
from typing import Optional
from pathlib import Path
from datetime import datetime
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

CTX_FILE = Path("/app/backend/context_store.json")

def _read_store():
    if not CTX_FILE.exists():
        return {}

    try:
        with open(CTX_FILE, "r") as f:
            return json.load(f)
    except Exception:
        return {}

def load_context(session_id, chat_id):
    data = _read_store()
    key = f"{session_id}:{chat_id}"
    return data.get(key, {})

class Filter:

    class Valves(BaseModel):
        priority: int = Field(default=-10)

    def __init__(self):
        self.valves = self.Valves()

    # =====================================================
    # --------- TABLE STYLING HELPERS FOR DOCX ------------
    # =====================================================

    def _shade_cell(self, cell, color="D9EAF7"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    def _set_table_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")

        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            borders.append(border)

        tblPr.append(borders)

    def _set_cell_border(self, cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = OxmlElement("w:tcBorders")

        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "6")  # slightly thicker
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            border.set(qn("w:themeColor"), "")  
            tcBorders.append(border)

        tcPr.append(tcBorders)

    def _set_cell_width(self, cell, width_mm):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:type"), "dxa")
        tcW.set(qn("w:w"), str(int(width_mm * 56.7)))
        tcPr.append(tcW)

    def _repeat_header_row(self, row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        tblHeader.set(qn("w:val"), "true")
        trPr.append(tblHeader)

    def _write_multiline_text(self, cell, text):
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        lines = text.split("\n")

        for i, line in enumerate(lines):
            run = paragraph.add_run(line)
            if i < len(lines) - 1:
                run.add_break()

    # =====================================================
    # ---------------- JSON EXTRACTION --------------------
    # =====================================================

    def _strip_fences(self, content: str) -> str:
        if content.startswith("```"):
            lines = content.split("\n")[1:]
            if lines and lines[-1].startswith("```"):
                lines = lines[:-1]
            return "\n".join(lines).strip()
        return content

    def _extract_json_block(self, text: str) -> str:
        match = re.search(r"\{[\s\S]*\}", text)
        return match.group(0) if match else ""

    # =====================================================
    # ---------------- TABLE RENDERING --------------------
    # =====================================================

    def _format_cell(self, text: str) -> str:
        if not text:
            return ""
        return text.replace("\n", "<br>")

    def _json_to_markdown_table(self, title: str, test_cases):

        headers = [
            "Step No",
            "Description",
            "Expected Results",
            "Actual Results",
            "Pass/Fail",
            "Comments",
        ]

        lines = []
        lines.append(f"### ✅ {title}\n")
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join(["---"] * len(headers)) + "|")

        for case in test_cases:
            row = [
                str(case.get("step_no", "")),
                case.get("description", ""),
                case.get("expected_results", ""),
                self._format_cell(case.get("actual_results", "")),
                self._format_cell(case.get("pass_fail", "")),
                case.get("comments", ""),
            ]

            row = [col.replace("|", "\\|") for col in row]
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    # =====================================================
    # ---------------- DOCX GENERATOR ---------------------
    # =====================================================

    def _generate_docx(self, data: dict) -> str:

        doc = Document()
        section = doc.sections[0]

        # Landscape
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )

        # Narrow margins
        section.top_margin = Mm(12.7)
        section.bottom_margin = Mm(12.7)
        section.left_margin = Mm(12.7)
        section.right_margin = Mm(12.7)

        # Global font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        document_title = data.get("document_title", "Validation Test Cases")

        # =====================================================
        # HEADER
        # =====================================================
        header = section.header

        # REMOVE default empty paragraph
        if header.paragraphs:
            p = header.paragraphs[0]
            p._element.getparent().remove(p._element)

        usable_width = section.page_width - section.left_margin - section.right_margin

        header_table = header.add_table(rows=3, cols=3, width=usable_width)
        header_table.autofit = False

        col_widths = [160, 50, 50]

        for row in header_table.rows:
            for cell in row.cells:
                self._set_cell_border(cell)

        for row in header_table.rows:
            for i, cell in enumerate(row.cells):
                self._set_cell_width(cell, col_widths[i])

        # Merge A1 + B1
        header_table.cell(0, 0).merge(header_table.cell(1, 0))
        logo_path = "/app/backend/static/logo.png"

        cell = header_table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()

        if Path(logo_path).exists():
            run.add_picture(logo_path, width=Inches(2))

        header_table.cell(2, 0).text = document_title
        header_table.cell(0, 1).text = "Document ID:"
        header_table.cell(1, 1).text = "Revision No:"
        header_table.cell(1, 2).text = "01.00"
        header_table.cell(2, 1).text = "Document Date:"
        header_table.cell(2, 2).text = "Refer ZenQMS"

        for row in header_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        header.add_paragraph("")

        # =====================================================
        # FOOTER
        # =====================================================
        footer = section.footer
        footer_table = footer.add_table(rows=1, cols=3, width=usable_width)

        cell_left = footer_table.cell(0, 0)
        cell_center = footer_table.cell(0, 1)
        cell_right = footer_table.cell(0, 2)

        cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        footer_table.cell(0, 0).text = f"File: {document_title}"
        footer_table.cell(0, 1).text = "PROPRIETARY & CONFIDENTIAL"

        cell_center.paragraphs[0].alignment = 1

        p = footer_table.cell(0, 2).paragraphs[0]
        p.text = "Page "
        run = p.add_run()

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"

        fldCharEnd = OxmlElement("w:fldChar")
        fldCharEnd.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldCharEnd)

        cell_right.paragraphs[0].alignment = 2

        # =====================================================
        # REVISION HISTORY PAGE
        # =====================================================
        p = doc.add_paragraph()
        p.alignment = 1
        run = p.add_run("Revision History")
        run.bold = True
        run.underline = True

        rev_table = doc.add_table(rows=1, cols=4)
        self._set_table_borders(rev_table)

        headers = ["Version", "Date", "Author", "Description"]
        widths = [30, 40, 50, 140]

        for i, text in enumerate(headers):
            cell = rev_table.rows[0].cells[i]
            cell.text = text
            self._set_cell_width(cell, widths[i])

        for cell in rev_table.rows[0].cells:
            self._shade_cell(cell, "D9EAF7")

        row = rev_table.add_row().cells
        row[0].text = "1.0"
        row[1].text = datetime.now().strftime("%d-%b-%Y")
        row[2].text = "System"
        row[3].text = "Initial Document Generation"

        for row in rev_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.add_page_break()

        # =====================================================
        # TEST SCRIPT TABLE
        # =====================================================
        doc.add_heading("Test Script", level=3)

        table = doc.add_table(rows=1, cols=6)
        table.autofit = False
        self._set_table_borders(table)

        headers = [
            "Step No",
            "Description",
            "Expected Results",
            "Actual Results",
            "Pass/Fail",
            "Comments",
        ]

        col_widths = [20, 70, 70, 30, 30, 40]

        for i, header_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header_text
            self._set_cell_width(cell, col_widths[i])

        for cell in table.rows[0].cells:
            self._shade_cell(cell, "BFBFBF")

        header_row = table.rows[0]

        for cell in header_row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        self._repeat_header_row(table.rows[0])

        for case in data["test_cases"]:
            row = table.add_row().cells

            row[0].text = str(case.get("step_no", ""))
            row[1].text = case.get("description", "")
            row[2].text = case.get("expected_results", "")

            self._write_multiline_text(
                row[3], case.get("actual_results", "☐ As Expected\n☐ Other")
            )
            self._write_multiline_text(row[4], case.get("pass_fail", "☐ Pass\n☐ Fail"))

            row[5].text = case.get("comments", "")

            for i, cell in enumerate(row):
                self._set_cell_width(cell, col_widths[i])

        # Save
        output_dir = Path("/app/backend/docx_output")
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = document_title.replace(" ", "_")
        file_path = output_dir / f"{safe_title}_{timestamp}.docx"

        doc.save(file_path)

        print(f"[DOCX] Saved at: {file_path}")

        return str(file_path)

    # =====================================================
    # ---------------- MAIN FILTER ------------------------
    # =====================================================

    def outlet(self, body: dict, __user__: Optional[dict] = None) -> dict:
        print("[UltraAssist OpenWebUI Filter] Format Output and Generate DOCX started.")
        try:

            messages = body.get("messages")
            if not messages:
                return body

            print(f"[UltraAssist OpenWebUI Filter] Body: {body}")

            session_id = body.get("session_id")
            chat_id = body.get("chat_id")

            ctx = load_context(session_id, chat_id)

            department = ctx.get("department")
            purpose = ctx.get("purpose")

            print(f"[UltraAssist OpenWebUI Filter] Extracted Department: {department}, Purpose: {purpose}")

            # Skip for all non-validation/script_authoring
            if not (department == "validation" and purpose == "script_authoring"):
                return body

            # Only process LAST assistant message
            for msg in reversed(messages):

                if msg.get("role") != "assistant":
                    continue

                content = msg.get("content", "").strip()
                if not content:
                    return body

                # Skip if already table
                if content.startswith("| Step No"):
                    return body

                content_clean = self._strip_fences(content)
                json_str = self._extract_json_block(content_clean)

                if not json_str:
                    return body

                data = json.loads(json_str)

                if "test_cases" not in data:
                    return body

                file_path = self._generate_docx(data)

                # -------- Build Table --------
                table_md = self._json_to_markdown_table(
                    data.get("document_title", "Validation Test Cases"),
                    data["test_cases"],
                )

                filename = Path(file_path).name

                footer_note = f"\n\n📄 **File Generated:** `{filename}`"

                msg["content"] = table_md + footer_note

                return body

            return body

        except Exception as e:
            print("FILTER ERROR:", str(e))
            return body
