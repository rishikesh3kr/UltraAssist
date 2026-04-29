"""
title: Download Generated DOCX
author: Rishikesh Kumar
version: 1.0
icon_url: data:image/svg+xml;base64,PD94bWwgdmVyc2lvbj0iMS4wIiBlbmNvZGluZz0idXRmLTgiPz48IS0tIFVwbG9hZGVkIHRvOiBTVkcgUmVwbywgd3d3LnN2Z3JlcG8uY29tLCBHZW5lcmF0b3I6IFNWRyBSZXBvIE1peGVyIFRvb2xzIC0tPgo8c3ZnIGZpbGw9IiMwMDAwMDAiIHdpZHRoPSI4MDBweCIgaGVpZ2h0PSI4MDBweCIgdmlld0JveD0iMCAwIDI0IDI0IiBpZD0iZG93bmxvYWQtYWx0IiBkYXRhLW5hbWU9IkZsYXQgTGluZSIgeG1sbnM9Imh0dHA6Ly93d3cudzMub3JnLzIwMDAvc3ZnIiBjbGFzcz0iaWNvbiBmbGF0LWxpbmUiPjxwb2x5bGluZSBpZD0icHJpbWFyeSIgcG9pbnRzPSI4IDEyIDEyIDE2IDE2IDEyIiBzdHlsZT0iZmlsbDogbm9uZTsgc3Ryb2tlOiByZ2IoMCwgMCwgMCk7IHN0cm9rZS1saW5lY2FwOiByb3VuZDsgc3Ryb2tlLWxpbmVqb2luOiByb3VuZDsgc3Ryb2tlLXdpZHRoOiAyOyI+PC9wb2x5bGluZT48cGF0aCBpZD0icHJpbWFyeS0yIiBkYXRhLW5hbWU9InByaW1hcnkiIGQ9Ik01LDIxSDE5TTEyLDNWMTYiIHN0eWxlPSJmaWxsOiBub25lOyBzdHJva2U6IHJnYigwLCAwLCAwKTsgc3Ryb2tlLWxpbmVjYXA6IHJvdW5kOyBzdHJva2UtbGluZWpvaW46IHJvdW5kOyBzdHJva2Utd2lkdGg6IDI7Ij48L3BhdGg+PC9zdmc+
"""

import json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydoc import doc
import re
from docx.shared import Pt
import base64
from pathlib import Path
from numpy.char import title
from pydantic import BaseModel, Field
from docx import Document
from datetime import datetime
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

CTX_FILE = Path("/app/backend/context_store.json")

def _read_store():
    if not CTX_FILE.exists():
        return {}

    try:
        with open(CTX_FILE, "r") as f:
            return json.load(f)
    except Exception:
        return {}

def load_context(session_id, chat_id):
    data = _read_store()
    key = f"{session_id}:{chat_id}"
    return data.get(key, {})

class Action:

    class Valves(BaseModel):
        priority: int = Field(default=0)

    def __init__(self):
        self.valves = self.Valves()

    def _sanitize_filename(self, text: str) -> str:
        # Replace spaces with underscore
        text = text.strip().replace(" ", "_")

        # Remove invalid characters
        text = re.sub(r'[\\/*?:"<>|]', '', text)

        # Remove multiple underscores
        text = re.sub(r'_+', '_', text)

        # Limit length (important)
        return text[:80] if len(text) > 80 else text

    def _add_footer(self, doc, title=""):
        section = doc.sections[0]
        footer = section.footer

        # Remove default paragraph
        if footer.paragraphs:
            p = footer.paragraphs[0]
            p._element.getparent().remove(p._element)

        # Create table (3 columns)
        usable_width = section.page_width - section.left_margin - section.right_margin

        table = footer.add_table(rows=1, cols=3, width=usable_width)
        table.autofit = True

        # -----------------------------------
        # LEFT → File Name
        # -----------------------------------
        cell_left = table.cell(0, 0)
        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_left.add_run(f"File: {title}")

        # -----------------------------------
        # CENTER → Confidential Text
        # -----------------------------------
        cell_center = table.cell(0, 1)
        p_center = cell_center.paragraphs[0]
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p_center.add_run("PROPRIETARY & CONFIDENTIAL")
        run.bold = True
        run.font.size = Pt(9)

        # -----------------------------------
        # RIGHT → Page Number
        # -----------------------------------
        cell_right = table.cell(0, 2)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p_right.add_run("Page ")

        run = p_right.add_run()

        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"

        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)

    def _add_header_logo(self, doc, title=""):

        section = doc.sections[0]
        header = section.header

        logo_path = "/app/backend/static/Logo.png"

        # Remove default paragraph
        if header.paragraphs:
            p = header.paragraphs[0]
            p._element.getparent().remove(p._element)

        usable_width = section.page_width - section.left_margin - section.right_margin

        table = header.add_table(rows=1, cols=2, width=usable_width)
        table.autofit = True

        # LEFT → LOGO
        cell_left = table.cell(0, 0)
        p_left = cell_left.paragraphs[0]

        if Path(logo_path).exists():
            run = p_left.add_run()
            run.add_picture(logo_path, width=Inches(1.5))

        # RIGHT → TITLE + DATE
        cell_right = table.cell(0, 1)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = p_right.add_run(title + "\n")
        run.bold = True

        p_right.add_run(datetime.now().strftime("%d-%b-%Y"))

    def _markdown_to_docx(self, doc, text: str):

        import re

        lines = text.split("\n")
        table_buffer = []
        in_table = False

        for line in lines:

            line = line.rstrip()

            # =========================
            # HEADINGS
            # =========================
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
                continue
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
                continue
            elif line.startswith("# "):
                continue  # already used as title

            # =========================
            # TABLE DETECTION
            # =========================
            if line.startswith("|") and "|" in line[1:]:
                table_buffer.append(line)
                in_table = True
                continue
            else:
                if in_table:
                    self._render_table(doc, table_buffer)
                    table_buffer = []
                    in_table = False

            # =========================
            # NUMBERED LIST
            # =========================
            if re.match(r"^\d+\.\s+", line):
                doc.add_paragraph(line, style="List Number")
                continue

            # =========================
            # BULLETS
            # =========================
            if line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
                continue

            # =========================
            # CODE BLOCK
            # =========================
            if line.startswith("```"):
                doc.add_paragraph("[Code Block]")
                continue

            # =========================
            # NORMAL TEXT WITH FORMATTING
            # =========================
            if line.strip():
                p = doc.add_paragraph()
                self._add_inline_formatting(p, line)

        if table_buffer:
            self._render_table(doc, table_buffer)

    def _set_cell_border(self, cell):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        tcBorders = OxmlElement('w:tcBorders')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)

        tcPr.append(tcBorders)

    def _shade_cell(self, cell, color="D9D9D9"):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def _render_table(self, doc, table_lines):

        section = doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin

        rows = [row.strip("|").split("|") for row in table_lines]

        # Remove markdown separator row
        if len(rows) > 1 and all("---" in cell for cell in rows[1]):
            rows.pop(1)

        cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=cols)
        table.autofit = False   

        col_width = usable_width // cols

        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = cell_text.strip()

                # Set column width
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:type"), "dxa")
                tcW.set(qn("w:w"), str(col_width))
                tcPr.append(tcW)

                self._set_cell_border(cell)

                # Header styling
                if i == 0:
                    self._shade_cell(cell, "BFBFBF")

    def _add_inline_formatting(self, paragraph, text):

        import re

        tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)

        for token in tokens:

            if token.startswith("**") and token.endswith("**"):
                run = paragraph.add_run(token[2:-2])
                run.bold = True

            elif token.startswith("*") and token.endswith("*"):
                run = paragraph.add_run(token[1:-1])
                run.italic = True

            else:
                paragraph.add_run(token)

    def _extract_title(self, text: str) -> str:
        for line in text.split("\n"):
            line = line.strip()
            if line.startswith("# "):
                return line.replace("# ", "").strip()
        return "Chatbot Response"

    def _extract_filename(self, content: str) -> str:
        match = re.search(r"`([^`]+\.docx)`", content)
        return match.group(1) if match else ""

    async def action(self, body: dict, __event_emitter__=None, __event_call__=None, **kwargs):
        try:
            if __event_emitter__:
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": "Preparing file for download...",
                            "done": False,
                        },
                    }
                )

            print(f"[UltraAssist OpenWebUI Filter] Received body: {body}")

            session_id = body.get("session_id")
            chat_id = body.get("chat_id")

            ctx = load_context(session_id, chat_id)

            department = ctx.get("department")
            purpose = ctx.get("purpose")

            print(f"[UltraAssist OpenWebUI Filter] Extracted Department: {department}, Purpose: {purpose}")

            messages = body.get("messages", [])
            filename = ""
            file_path = None

            # =====================================================
            # 🟢 CASE 1: Validation Script Authoring
            # =====================================================
            if department == "validation" and purpose == "script_authoring":

                for msg in reversed(messages):
                    if msg.get("role") == "assistant":
                        filename = self._extract_filename(msg.get("content", ""))
                        break

                if not filename:
                    return {"content": "❌ Could not detect file name."}

                file_path = Path("/app/backend/docx_output") / filename

                if not file_path.exists():
                    return {"content": f"❌ File not found: {filename}"}

            # =====================================================
            # 🔵 CASE 2: Generic Queries → Generate DOCX from text
            # =====================================================
            else:

                assistant_text = ""

                for msg in reversed(messages):
                    if msg.get("role") == "assistant":
                        assistant_text = msg.get("content", "")
                        break

                if not assistant_text:
                    return {"content": "❌ No assistant response found."}

                doc = Document()

                # Global style
                style = doc.styles["Normal"]
                style.font.name = "Calibri"
                style.font.size = Pt(11)

                # Title
                title = self._extract_title(assistant_text)
                
                #Add header with logo, footer and title
                self._add_header_logo(doc, title)
                self._add_footer(doc, title)
                doc.add_heading(title, level=1)

                # Content
                self._markdown_to_docx(doc, assistant_text)

                output_dir = Path("/app/backend/docx_output")
                output_dir.mkdir(parents=True, exist_ok=True)

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_title = self._sanitize_filename(title)
                filename = f"{safe_title[:30]}_{timestamp}.docx"

                file_path = output_dir / filename

                doc.save(file_path)

            # Read file and encode
            with open(file_path, "rb") as f:
                file_bytes = f.read()

            file_base64 = base64.b64encode(file_bytes).decode("utf-8")

            # Trigger browser download
            if __event_call__:
                await __event_call__(
                    {
                        "type": "execute",
                        "data": {
                            "code": f"""
try {{
    const base64Data = "{file_base64}";
    const binaryData = atob(base64Data);
    const arrayBuffer = new Uint8Array(binaryData.length);
    for (let i = 0; i < binaryData.length; i++) {{
        arrayBuffer[i] = binaryData.charCodeAt(i);
    }}
    const blob = new Blob([arrayBuffer], {{
        type: "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }});
    const url = URL.createObjectURL(blob);
    const a = document.createElement("a");
    a.style.display = "none";
    a.href = url;
    a.download = "{filename}";
    document.body.appendChild(a);
    a.click();
    URL.revokeObjectURL(url);
    document.body.removeChild(a);
}} catch (error) {{
    console.error("Download failed:", error);
}}
"""
                        },
                    }
                )

            if __event_emitter__:
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {"description": "Download started.", "done": True},
                    }
                )

            return {"content": "📄 Download started."}

        except Exception as e:
            if __event_emitter__:
                await __event_emitter__(
                    {
                        "type": "notification",
                        "data": {"type": "error", "content": str(e)},
                    }
                )

            return {"content": f"❌ Error: {str(e)}"}
