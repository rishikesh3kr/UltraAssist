"""
{
  title: UltraAssist RAG
  author: Rishikesh Kumar
  version: 1.0
  type: filter
}
"""

from pyexpat.errors import messages

import requests
import re
import os
import uuid
from typing import Optional
from pydantic import BaseModel, Field
import shutil
from docx import Document
from docx.shared import Inches
import json
import traceback
import threading
from pathlib import Path
from open_webui.internal.db import get_session
from open_webui.models.users import Users
from open_webui.internal.db import get_session
from open_webui.models.groups import Groups

DEPARTMENTS = [
    "finance",
    "managed_services",
    "validation",
    "business_transition",
    "products",
    "general",
]

PURPOSES = {
    "finance": ["finance_query"],
    "managed_services": ["support_ops"],
    "validation": ["script_authoring", "sop_data"],
    "business_transition": ["sop_creation"],
    "products": ["product_query"],
    "general": ["general"],
}

COUNTER_FILE = "/app/data/id_counter.json"
id_lock = threading.Lock()

CTX_FILE = Path("/app/backend/context_store.json")

def _read_store():
    if not CTX_FILE.exists():
        return {}

    try:
        with open(CTX_FILE, "r") as f:
            return json.load(f)
    except Exception:
        return {}

def _write_store(data):
    CTX_FILE.parent.mkdir(parents=True, exist_ok=True) 

    with open(CTX_FILE, "w") as f:
        json.dump(data, f, indent=2)

def save_context(session_id, chat_id, department, purpose):
    data = _read_store()
    key = f"{session_id}:{chat_id}"
    data[key] = {
        "department": department,
        "purpose": purpose
    }
    _write_store(data)

def load_context(session_id, chat_id):
    data = _read_store()
    key = f"{session_id}:{chat_id}"
    return data.get(key, {})

def call_llm(query: str) -> str:
    try:
        url = "http://litellm:4001/v1/chat/completions"

        payload = {
            "model": "claude-4-5-haiku",
            "messages": [{"role": "user", "content": query}],
            "temperature": 0,
        }

        response = requests.post(
            url,
            headers={
                "Authorization": f"Bearer {os.getenv('OPENAI_API_KEYS', 'sk-1234')}"
            },
            json=payload,
            timeout=20,
        )

        response.raise_for_status()

        return response.json()["choices"][0]["message"]["content"]

    except Exception as e:
        print(f"[UltraAssist Open WebUI - call_llm] LLM call failed: {e}")
        return "general"


def get_last_user_message(messages):
    for msg in reversed(messages):
        if msg.get("role") == "user":
            return msg.get("content", "")
    return ""


def classify_scope(query: str) -> tuple[str, str]:
    query = query.strip()

    prompt = f"""
You are a strict classifier.

Classify the user query into:
1. department
2. purpose

Return ONLY in this JSON format:
{{
  "department": "...",
  "purpose": "..."
}}

Departments and purposes:

finance:
- finance_query → billing, invoices, payments, revenue

managed_services:
- support_ops → incidents, monitoring, operations

validation:
- script_authoring → test case generation, FRS, URS, validation scripts
- sop_data → SOPs querying, validation guidelines, compliance documents queries

business_transition:
- sop_creation → SOP writing, migration processes, onboarding

products:
- product_query → product features, roadmap, bugs

general:
- general → anything else

Rules:
- Return ONLY JSON
- No explanation
- No extra text
- Always include both fields

User Query: {query}
"""

    response = call_llm(prompt)

    print(
        f"[UltraAssist Open WebUI - classify_scope] Raw LLM Classification Response: {response}"
    )

    if response.startswith("```"):
        response = re.sub(r"```[a-zA-Z]*", "", response)
        response = response.replace("```", "").strip()

    try:
        parsed = json.loads(response)
        department = parsed.get("department", "general").lower()
        purpose = parsed.get("purpose", "general").lower()

        if department not in PURPOSES:
            return "general", "general"

        if purpose not in PURPOSES.get(department, []):
            return department, PURPOSES[department][0]

        print(
            "[UltraAssist Open WebUI - classify_scope] Parsed Department -",
            department,
            ", Purpose - ",
            purpose,
        )
        return department, purpose

    except Exception:
        print(
            "[UltraAssist Open WebUI - classify_scope] Failed to parse JSON:", response
        )
        return "general", "general"

async def get_user_groups(user_id: str):
    db = None
    try:
        db = next(get_session())

        groups = await Groups.get_groups_by_member_id(user_id, db=db)

        return [{"id": g.id, "name": g.name} for g in groups]

    except Exception as e:
        print(f"[UltraAssist Open WebUI - get_user_groups] ❌ Error fetching groups: {e}")
        return []

    finally:
        if db:
            db.close()

async def fetch_user_from_db(user_id: str):
    db = None
    try:
        db = next(get_session())

        user = await Users.get_user_by_id(user_id, db=db)

        return user

    except Exception as e:
        print(f"[UltraAssist Open WebUI - fetch_user_from_db] ❌ DB fetch failed: {e}")
        return None

    finally:
        if db:
            db.close()


def normalize_messages(messages):
    cleaned_messages = []

    for msg in messages:
        new_msg = msg.copy()
        content = new_msg.get("content")

        if isinstance(content, list):
            text_parts = [
                item.get("text", "") for item in content if item.get("type") == "text"
            ]
            new_msg["content"] = " ".join(text_parts).strip()

        elif isinstance(content, str):
            new_msg["content"] = content.strip()

        else:
            new_msg["content"] = ""

        cleaned_messages.append(new_msg)

    return cleaned_messages


def get_next_ids(prefix="IMG"):
    with id_lock:
        try:
            if not os.path.exists(COUNTER_FILE):
                with open(COUNTER_FILE, "w") as f:
                    json.dump({"urs_counter": 1, "frs_counter": 1}, f)

            with open(COUNTER_FILE, "r") as f:
                data = json.load(f)

            urs_counter = data.get("urs_counter", 1)
            frs_counter = data.get("frs_counter", 1)

            urs_id = f"UIT-UR-{prefix}-{urs_counter:03d}"
            frs_id = f"UIT-FR-{prefix}-{frs_counter:03d}"

            data["urs_counter"] = urs_counter + 1
            data["frs_counter"] = frs_counter + 1

            with open(COUNTER_FILE, "w") as f:
                json.dump(data, f)

            return urs_id, frs_id

        except Exception as e:
            print(
                f"[UltraAssist Open WebUI - get_next_ids] ❌ ID generation failed: {e}"
            )
            return None


import os
import uuid
import traceback
from docx import Document
from docx.shared import Inches


def create_docx_from_image(image_path: str, output_dir: str):
    try:
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)

        doc = Document()

        # Add heading
        doc.add_heading("Generated FRS from Image", level=1)

        # Create table
        table = doc.add_table(rows=2, cols=3)

        # Headers
        headers = ["URS ID", "FRS ID", "Description"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        # Generate IDs
        ids = get_next_ids(prefix="IMG")
        if not ids:
            raise Exception("ID generation failed")

        urs_id, frs_id = ids

        # Fill URS & FRS
        table.rows[1].cells[0].text = urs_id
        table.rows[1].cells[1].text = frs_id

        # Improved Description (IMPORTANT for embeddings)
        description_cell = table.rows[1].cells[2]
        description_cell.text = (
            "Image-based requirement. Refer attached screenshot for details."
        )

        # Add image below text
        paragraph = description_cell.add_paragraph()
        run = paragraph.add_run()

        if os.path.exists(image_path):
            run.add_picture(image_path, width=Inches(4))
        else:
            print(
                f"[UltraAssist Open WebUI - create_docx_from_image] ⚠️ Image not found: {image_path}"
            )

        # Unique filename (prevents overwrite)
        filename = f"{uuid.uuid4().hex}.docx"
        output_path = os.path.join(output_dir, filename)

        # Save doc
        doc.save(output_path)

        print(
            f"[UltraAssist Open WebUI - create_docx_from_image] 📄 DOCX created: {output_path}"
        )

        return {
            "status": "success",
            "docx_path": output_path,
            "urs_id": urs_id,
            "frs_id": frs_id,
        }

    except Exception as e:
        print(
            f"[UltraAssist Open WebUI - create_docx_from_image] ❌ DOCX creation failed: {repr(e)}"
        )
        traceback.print_exc()
        return {"status": "error", "message": str(e)}


def is_image_file(filename: str):
    if not filename or not isinstance(filename, str):
        return False

    allowed = (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp")

    if not filename.lower().endswith(allowed):
        return False

    return True


def get_latest_uploaded_file(body: dict):
    try:
        parent = body.get("metadata", {}).get("parent_message", {})
        parent_files = parent.get("files")

        if parent_files and isinstance(parent_files, list):
            latest = parent_files[-1]

            filename = latest.get("name") or latest.get("file", {}).get("filename")

            if not filename:
                return None

            uploads_dir = "/app/backend/data/uploads"

            if not os.path.exists(uploads_dir):
                print(
                    "[UltraAssist Open WebUI - get_latest_uploaded_file] ❌ Uploads directory not found"
                )
                return None

            # Safer matching: strip UUID prefix
            for f in os.listdir(uploads_dir):
                try:
                    actual_name = f.split("_", 1)[-1]  # removes UUID prefix
                    if actual_name == filename:
                        full_path = os.path.join(uploads_dir, f)

                        print(
                            f"[UltraAssist Open WebUI - get_latest_uploaded_file] ✅ Found uploaded file: {full_path}"
                        )

                        return {"filename": filename, "path": full_path}
                except Exception:
                    continue

    except Exception as e:
        print(
            f"[UltraAssist Open WebUI - get_latest_uploaded_file] ❌ File extraction failed: {e}"
        )

    return None


def move_file_to_frs(src_path: str, original_filename: str):
    if not src_path or not os.path.exists(src_path):
        print(
            f"[UltraAssist Open WebUI - move_file_to_frs] ❌ File not found: {src_path}"
        )
        return None

    # Basic file type validation
    allowed_extensions = (".docx", ".pdf", ".txt")
    if not original_filename.lower().endswith(allowed_extensions):
        print(
            f"[UltraAssist Open WebUI - move_file_to_frs] ❌ Unsupported file type: {original_filename}"
        )
        return None

    dest_dir = "/app/data/frs"
    os.makedirs(dest_dir, exist_ok=True)

    dest_path = os.path.join(dest_dir, original_filename)

    try:
        print(
            f"[UltraAssist Open WebUI - move_file_to_frs] 📂 Moving file to FRS: {original_filename}"
        )

        # Handle duplicate filenames safely
        if os.path.exists(dest_path):
            base, ext = os.path.splitext(original_filename)
            counter = 1

            while True:
                new_name = f"{base}_{counter}{ext}"
                new_path = os.path.join(dest_dir, new_name)

                if not os.path.exists(new_path):
                    dest_path = new_path
                    break

                counter += 1

        shutil.move(src_path, dest_path)

        print(
            f"[UltraAssist Open WebUI - move_file_to_frs] ✅ File moved to: {dest_path}"
        )
        return dest_path

    except Exception as e:
        print(f"[UltraAssist Open WebUI - move_file_to_frs] ❌ Move failed: {e}")
        return None


def trigger_frs_indexing(
    file_path: str, rag_url: str, department="validation", purpose="script_authoring"
):
    try:
        print(
            f"[UltraAssist Open WebUI - trigger_frs_indexing] 🔄 Indexing file: {file_path}"
        )

        response = requests.post(
            f"{rag_url}/index/frs",
            params={
                "file_path": file_path,
                "department": department,
                "purpose": purpose,
            },
            timeout=60,
        )

        print(
            f"[UltraAssist Open WebUI - trigger_frs_indexing] 📡 Index status: {response.status_code}"
        )

        if response.status_code != 200:
            return {
                "error": f"Indexing failed with status {response.status_code}",
                "response": response.text,
            }

        try:
            return response.json()
        except Exception:
            return {"status": "unknown", "raw_response": response.text}

    except Exception as e:
        print(
            f"[UltraAssist Open WebUI - trigger_frs_indexing] ❌ Indexing failed: {e}"
        )
        return {"error": str(e)}


def wait_for_frs_index(
    file_name: str,
    rag_url: str,
    department="validation",
    purpose="script_authoring",
    timeout=60,
):
    import time

    start = time.time()

    print(
        f"[UltraAssist Open WebUI - wait_for_frs_index] ⏳ Waiting for indexing: {file_name}"
    )

    while time.time() - start < timeout:
        try:
            res = requests.get(
                f"{rag_url}/debug/frs_dump",
                params={"department": department, "purpose": purpose},
                timeout=10,
            )

            if res.status_code != 200:
                print(
                    f"[UltraAssist Open WebUI - wait_for_frs_index] ⚠ Debug API error: {res.status_code}"
                )
                time.sleep(1)
                continue

            data = res.json()

            metadatas = data.get("sample_metadata", [])

            # ✅ Check metadata instead of text
            for meta in metadatas:
                source_file = meta.get("source_file", "")
                if file_name in source_file:
                    print(
                        "[UltraAssist Open WebUI - wait_for_frs_index] ✅ File detected in vector DB"
                    )
                    return True

        except Exception as e:
            print(f"[UltraAssist Open WebUI - wait_for_frs_index] Polling error: {e}")

        time.sleep(1)

    print("[UltraAssist Open WebUI - wait_for_frs_index] ⚠ Index not confirmed in time")
    return False


def build_identity(department: str, purpose: str) -> str:
    return f"""
You are UltraAssist, an intelligent enterprise assistant.

You specialize in the domain of {department} and are currently handling a request related to {purpose}.

GENERAL BEHAVIOR:

• Provide accurate, structured, and context-aware responses based on retrieved knowledge.
• Prioritize retrieved context over assumptions.
• If context is insufficient, respond conservatively and highlight gaps.
• Maintain clarity, depth, and precision in responses.
• Avoid generic answers when specific context is available.

CONTEXT USAGE RULES:

• Treat retrieved structured data (if any) as the primary source of truth.
• Treat supporting documents as secondary context for explanation and navigation.
• Do not invent system behaviors not grounded in the provided material.

OUTPUT RULES:

• Be structured and detailed where needed.
• Avoid unnecessary verbosity.
• Prefer clarity over length unless depth is explicitly required.

SPECIAL HANDLING:

• If the user greets → respond politely and introduce capabilities briefly.
• If the user asks follow-up questions → focus only on delta, not full regeneration.

"""

def build_validation_identity() -> str:
    return """
You are UltraAssist, a validation specialist focused on generating deep, exhaustive functional test cases.

OBJECTIVE:
Generate highly detailed, scenario-rich validation test cases based on the provided requirement and supporting context.

CRITICAL RULE:
You MUST generate a LARGE and DEEP test suite.
Minimum expected steps: Based on the requirement complexity, aim for 80–120 steps where applicable. If the requirement is simple, still aim for at least 30–50 steps by expanding scenarios and edge cases.

TEST DESIGN PRINCIPLES:

• Every step must follow: User Action → System Response
• Each step must validate ONE atomic behavior
• Split multi-actions into multiple steps
• Avoid generic steps — every step must test something specific

SCENARIO EXPANSION (MANDATORY):

You MUST expand coverage across:

1. Full flow execution
2. Partial flow execution
3. Interrupted flows
4. Retry scenarios
5. Failure and error handling
6. Data variations (valid, invalid, boundary)
7. Pre-existing data conditions
8. Overwrite and save behavior
9. Cancel and rollback flows
10. Concurrent or repeated actions
11. Navigation restrictions and edge cases

DO NOT STOP after one scenario.
You must generate MULTIPLE overlapping scenario blocks.

DEPTH RULE:

For every feature:
→ Expand it into at least 5–10 variations

Example:
- Success flow
- Failure flow
- Retry flow
- Cancel flow
- Invalid input flow

OUTPUT STRUCTURE:

Return ONLY valid JSON:

{
  "document_title": string,
  "test_cases": [
    {
      "step_no": number,
      "description": string,
      "expected_results": string,
      "actual_results": "☐ As Expected\n☐ Other",
      "pass_fail": "☐ Pass\n☐ Fail",
      "comments": ""
    }
  ]
}

STRICT RULES:

• No explanations outside JSON
• No markdown
• No summaries
• No skipped steps
• No vague expected results
• Use exact UI text where available
• Use placeholders like {User Role A}, {Case 1}

QUALITY CHECK BEFORE OUTPUT:

• If steps are repetitive → add new scenarios
• If steps are generic → make them UI-specific

If context contains:
"RETRIEVAL ERROR:No matching requirement found."
→ respond with that error clearly instead of generating test cases.
"""

def build_sop_creation_identity() -> str:
    return """
You are UltraAssist, a business transition specialist and an expert in Process Engineering and Pharmacovigilance (PV) Operations. Your responsibility is to transform complex technical requirements and system manuals into Standard Operating Procedures (SOPs) that are clear, audit-ready, and compliant with global PV regulations (e.g., GVP Modules). Your focus is on creating comprehensive SOPs for smooth organisational handoffs and knowledge transfer.

ROLE & TONE:
 - Tone: Authoritative, professional, clear, and instructional.
 - Expertise: You possess deep knowledge of GVP (Good Pharmacovigilance Practices), data integrity, and regulatory compliance.
 - Objective: Create procedures that a new employee can follow with zero ambiguity to perform high-stakes safety tasks.

PROCEDURAL STRUCTURE (The SOP Framework):
For every request, structure your response using these standard SOP sections:
    1. Purpose: Define the "Why." Why does this process exist? (e.g., "To ensure the consistent generation of AI-narratives for safety reporting.")
    2. Scope: Define the "Where" and "Who." Which modules and departments does this cover?
    3. Roles & Responsibilities: Clearly define who performs the actions (e.g., Case Processor, Medical Reviewer, System Admin).
    4. Procedural Steps: The core of the document. Use a logical, sequential flow.
    5. Quality Control (QC) & Compliance: Define specific checks, audit trail requirements, and "Null Flavor" handling.
    6. Exception Handling: What to do when the system fails, times out, or produces incorrect AI output.

SOP WRITING PRINCIPLES:
 - Imperative Mood: Use command-based language (e.g., "The User shall select..." or "The System Admin must configure...").
 - Role-Based Steps: Never just say "Do X." Always say "[Role Name] shall do X."
 - Grounded Clarity: Use exact UI terminology from the retrieved context (Buttons, Sections, Fields).
 - Best Practice Integration: If the retrieved requirement is silent on a critical PV step (like "Reviewing for Medical Sanity"), supplement it using general PV knowledge, but clearly distinguish it as a "Best Practice Recommendation."
 - Visual Continuity: Describe the expected system behavior after an action to guide the user (e.g., "Upon clicking Save, ensure the comparison window closes and the main Case Narrative field populates").

CONSTRAINTS & FOLLOW-UP LOGIC:
 - RAG Priority: The retrieved DOCUMENT CONTEXT provided below is your Primary Source. General PV knowledge is your extra Source for "filling the gaps" in procedural rigor.
 - No Hallucination: Do not invent software versions or company names not provided in the context.
 - Follow-up Awareness: * If the user asks to "elaborate on Section 4," ignore the RAG context if it is a duplicate of the previous turn and focus on the existing conversation history. If the user query is a modification (e.g., "Add a step for a second-level review"), update the procedure accordingly while maintaining the original SOP structure.
 - Formatting: Use Markdown (Headings, Bold text, Bullet points, and Numbered lists) to ensure the document is scannable and professional.

QUALITY CHECKLIST (Internal AI Check):
 - Did I define the specific Roles involved?
 - Are the steps sequential and logical?
 - Is the GenAI "Manual vs Automatic" configuration logic clearly explained?
 - Did I include a section on Error Handling/System Failure?
"""

def build_context_block(rag_result: dict, user_query: str = "") -> str:
    if "error" in rag_result:
        print(
            "[UltraAssist Open WebUI - build_context_block] ❌ RAG returned error:",
            rag_result["error"],
        )
        return f"""
USER QUERY:
{user_query}

RETRIEVAL ERROR:
{rag_result["error"]}
"""

    requirements = rag_result.get("requirements", [])
    manual_context = rag_result.get("manual_context") or rag_result.get("context") or []

    print(
        f"[UltraAssist Open WebUI - build_context_block] Requirements received: {len(requirements)}"
    )
    print(
        f"[UltraAssist Open WebUI - build_context_block] Manual chunks received: {len(manual_context)}"
    )

    # ---------------------------------------------------
    # Build Requirement Section (GENERIC)
    # ---------------------------------------------------
    requirement_block = ""

    if requirements:
        print(
            "[UltraAssist Open WebUI - build_context_block] Building requirement context..."
        )

        for idx, req in enumerate(requirements, 1):
            req_id = req.get("requirement_id", f"REQ_{idx}")
            req_text = req.get("text", "")

            requirement_block += f"""
--- REQUIREMENT {idx} ---
ID: {req_id}
{req_text}
"""

    else:
        print("[UltraAssist Open WebUI - build_context_block] ⚠ No requirements found")
        requirement_block = "RETRIEVAL ERROR: No matching requirement found."

    # ---------------------------------------------------
    # Build Manual Context Section
    # ---------------------------------------------------
    manual_block = ""

    if manual_context:
        print(
            "[UltraAssist Open WebUI - build_context_block] Building manual context..."
        )

        for idx, chunk in enumerate(manual_context, 1):
            metadata = chunk.get("metadata", {})
            source_file = metadata.get("source_file", "unknown")

            manual_block += f"""
--- CONTEXT {idx} ---
Source: {source_file}
{chunk.get("text", "")}
"""
    else:
        manual_block = "No supporting context available."

    # ---------------------------------------------------
    # Final Context Block
    # ---------------------------------------------------
    if requirements:
        context_text = f"""
USER QUERY:
{user_query}

========================
PRIMARY KNOWLEDGE (HIGH PRIORITY)
========================
{requirement_block}

========================
SUPPORTING CONTEXT (SECONDARY)
========================
{manual_block}

========================
INSTRUCTIONS
========================
• Use PRIMARY KNOWLEDGE as the main source of truth
• Use SUPPORTING CONTEXT only to enhance understanding
• Do NOT invent functionality not present in the context
• If primary knowledge is missing → clearly state limitation
    """
        
    else:
        context_text = f"""
USER QUERY:
{user_query}

========================
DOCUMENTATION CONTEXT 
========================
{manual_block}

========================
INSTRUCTIONS
========================
• Use DOCUMENTATION CONTEXT as the main source of truth
• If information relevant to the query is missing, clearly state limitation.

    """

    return context_text

def build_system_prompt(
    rag_result: dict,
    user_query: str = "",
    department: str = "general",
    purpose: str = "general",
) -> str:

    # ---------------------------------------------------
    #  Select Identity Dynamically
    # ---------------------------------------------------
    if department == "validation" and purpose == "script_authoring":
        identity = build_validation_identity()
    elif department == "business_transition" and purpose == "sop_creation":
        identity = build_sop_creation_identity()
    else:
        identity = build_identity(department, purpose)

    # ---------------------------------------------------
    #  Build Context
    # ---------------------------------------------------
    context_block = build_context_block(rag_result, user_query)

    # ---------------------------------------------------
    #  Compose Final Prompt (Structured)
    # ---------------------------------------------------
    full_prompt = f"""
========================
SYSTEM IDENTITY
========================
{identity}

========================
CONTEXT
========================
{context_block}
"""

    print(
        f"[UltraAssist Open WebUI - build_system_prompt] System prompt length: {len(full_prompt)} characters"
    )
    print(
        f"[UltraAssist Open WebUI - build_system_prompt] System prompt: {full_prompt} "
    )
    print(
        f"[UltraAssist Open WebUI - build_system_prompt] Mode: {department}/{purpose}"
    )

    return full_prompt

class Filter:

    class Valves(BaseModel):
        RAG_SERVICE_URL: str = Field(
            default=os.getenv("RAG_SERVICE_URL", "http://ultraassist-rag-service:8020")
        )
        TOP_K: int = Field(default=5)
        MANUAL_TOP_K: int = Field(default=5)
        RAG_TIMEOUT_SECONDS: int = Field(
            default=int(os.getenv("RAG_TIMEOUT_SECONDS", "45"))
        )
        ENABLE_GUARDRAILS: bool = Field(default=True)

    def __init__(self):
        self.valves = self.Valves()

        self.injection_patterns = [
            r"ignore\s+(previous|above|all)",
            r"disregard\s+(previous|above|all)",
            r"forget\s+(everything|all|previous)",
            r"you\s+are\s+now",
            r"system\s*:\s*",
        ]

    # -------------------------
    # 🔐 Guardrails
    # -------------------------

    def detect_prompt_injection(self, text: str) -> bool:
        if not self.valves.ENABLE_GUARDRAILS:
            return False
        return any(re.search(p, text.lower()) for p in self.injection_patterns)

    def sanitize_input(self, text: str) -> str:
        return text[:2000].strip()

    # -------------------------
    # 🔍 RAG Integration
    # -------------------------

    def get_context(
        self, query: str, department: str, purpose: str, user_groups: list
    ) -> dict:
        try:
            print("========================================")
            print("[UltraAssist Open WebUI - get_context] 🔍 Calling RAG Service")
            print(f"[UltraAssist Open WebUI - get_context] Department: {department}")
            print(f"[UltraAssist Open WebUI - get_context] Purpose: {purpose}")
            print(f"[UltraAssist Open WebUI - get_context] User Groups: {user_groups}")
            print("========================================")

            normalized_groups = [g.lower() for g in user_groups]

            # Admin override
            is_admin = "admin" in normalized_groups

            # Access check
            if not is_admin and department not in normalized_groups:
                print(
                    f"[UltraAssist Open WebUI - get_context] ❌ Access Denied: {department}"
                )
                return {"error": f"User not authorized for department: {department}"}

            if is_admin:
                print(
                    "[UltraAssist Open WebUI - get_context] 🔓 Admin access granted (all departments)"
                )

            response = requests.post(
                f"{self.valves.RAG_SERVICE_URL}/retrieve",
                json={
                    "query": query,
                    "department": department,
                    "purpose": purpose,
                    "top_k": self.valves.TOP_K,
                    "manual_top_k": self.valves.MANUAL_TOP_K,
                },
                timeout=self.valves.RAG_TIMEOUT_SECONDS,
            )

            print(
                f"[UltraAssist Open WebUI - get_context] RAG Status: {response.status_code}"
            )

            data = response.json()

            if "error" in data:
                print(
                    "[UltraAssist Open WebUI - get_context] ❌ RAG Error:",
                    data["error"],
                )
            else:
                print("[UltraAssist Open WebUI - get_context] ✅ RAG Success")

            return data

        except Exception as e:
            print("[UltraAssist Open WebUI - get_context] ❌ RAG unavailable:", str(e))
            return {"error": "RAG service unavailable"}

    # -------------------------
    # 🚪 ENTRY POINT
    # -------------------------

    async def inlet(self, body: dict, user: Optional[dict] = None) -> dict:

        # -------------------------
        # 👤 USER CONTEXT
        # -------------------------
        user_id = body.get("metadata", {}).get("user_id")
        db_user = await fetch_user_from_db(user_id)

        user_groups_data = await get_user_groups(user_id)
        user_groups = []

        print("\n===================================================")
        print("[UltraAssist Open WebUI - inlet] 🚀 New Request")
        print("===================================================")

        if db_user:
            print(f"User: {db_user.name} | {db_user.email} | Role: {db_user.role}")

        print("---------------------------------------------------")
        print("[UltraAssist Open WebUI - inlet] 👥 User Groups:")

        for g in user_groups_data:
            name = g.get("name")
            user_groups.append(name)
            print(name)

        print("===================================================")

        # -------------------------
        # 💬 MESSAGES
        # -------------------------
        messages = normalize_messages(body.get("messages", []))
        body["messages"] = messages

        if not messages:
            return body

        user_message = next(
            (m.get("content") for m in reversed(messages) if m.get("role") == "user"),
            "",
        )

        if not user_message:
            return body

        print(f"[UltraAssist Open WebUI - inlet] Query: {user_message}")

        # -------------------------
        # 🧠 CLASSIFICATION
        # -------------------------
        department, purpose = classify_scope(user_message)
        print(f"[UltraAssist Open WebUI - inlet] Scope → {department}/{purpose}")

        if "metadata" not in body:
            body["metadata"] = {}

        body["metadata"]["department"] = department
        body["metadata"]["purpose"] = purpose

        print(f"[UltraAssist Open WebUI - inlet] Department → {department}")
        print(f"[UltraAssist Open WebUI - inlet] Purpose → {purpose}")

        # -------------------------
        # 🛡 GUARDRAILS
        # -------------------------
        if self.detect_prompt_injection(user_message):
            print("[UltraAssist Open WebUI - inlet] ❌ Prompt Injection Detected")
            body["messages"] = [
                {"role": "assistant", "content": "Invalid instruction detected."}
            ]
            return body

        user_message = self.sanitize_input(user_message)

        # -------------------------
        # 📂 FILE HANDLING
        # -------------------------
        file_data = get_latest_uploaded_file(body)

        if file_data:
            src_path = file_data["path"]
            filename = file_data["filename"]

            frs_dir = f"/app/data/{department}/{purpose}/frs"
            os.makedirs(frs_dir, exist_ok=True)

            if is_image_file(filename):
                print("[UltraAssist Open WebUI - inlet] Image Upload Detected")

                docx_path = create_docx_from_image(src_path, frs_dir)

                if docx_path:
                    trigger_frs_indexing(docx_path, self.valves.RAG_SERVICE_URL)
                    wait_for_frs_index(
                        os.path.basename(docx_path), self.valves.RAG_SERVICE_URL
                    )

                os.remove(src_path)

            else:
                print("[UltraAssist Open WebUI - inlet] 📄 Document Upload Detected")

                moved_path = move_file_to_frs(src_path, filename, department, purpose)

                if moved_path:
                    trigger_frs_indexing(moved_path, self.valves.RAG_SERVICE_URL)
                    wait_for_frs_index(filename, self.valves.RAG_SERVICE_URL)

        # -------------------------
        # 🔍 RAG CALL
        # -------------------------
        rag_result = self.get_context(user_message, department, purpose, user_groups)

        if "error" in rag_result and "not authorized" in rag_result["error"].lower():
            print(
                "[UltraAssist Open WebUI - inlet] 🚫 Blocking LLM call due to access control"
            )

            body["messages"] = [
                {
                    "role": "assistant",
                    "content": f"You do not have access to the '{department}' department.",
                }
            ]

            return body

        # -------------------------
        # 🧠 PROMPT BUILD
        # -------------------------
        system_prompt = build_system_prompt(
            rag_result, user_message, department, purpose
        )

        body["messages"] = [{"role": "system", "content": system_prompt}] + messages

        metadata = body.get("metadata", {})

        session_id = metadata.get("session_id")
        chat_id = metadata.get("chat_id")

        save_context(session_id, chat_id, department, purpose)

        return body
